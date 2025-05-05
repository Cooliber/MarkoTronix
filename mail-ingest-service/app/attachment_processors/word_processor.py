"""
Word Document Attachment Processor

This module provides functionality for processing Word document attachments,
including text extraction and metadata extraction.
"""

import os
import re
import asyncio
from typing import Dict, List, Any, Tuple
from pathlib import Path

from .base import AttachmentProcessor, AttachmentProcessingResult

# Import optional dependencies - these will be checked at runtime
try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class WordProcessor(AttachmentProcessor):
    """Processor for Word document attachments."""
    
    def __init__(self):
        """Initialize the Word processor."""
        super().__init__()
        
        # Check if dependencies are available
        if not DOCX_AVAILABLE:
            raise ImportError(
                "Word processing dependencies not available. "
                "Please install python-docx."
            )
    
    @staticmethod
    def can_process(file_path: str) -> bool:
        """Check if this processor can handle the given file."""
        _, ext = os.path.splitext(file_path)
        return ext.lower() in ['.docx', '.doc']
    
    @staticmethod
    def get_mime_types() -> List[str]:
        """Get the MIME types this processor can handle."""
        return [
            'application/msword',
            'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        ]
    
    async def process(self, file_path: str) -> AttachmentProcessingResult:
        """Process a Word document attachment."""
        result = AttachmentProcessingResult()
        
        try:
            # Extract text and metadata
            text, metadata = await self._extract_text_and_metadata(file_path)
            
            # Extract entities from the text
            entities = await self._extract_entities(text)
            
            # Determine tags based on content
            tags = await self._determine_tags(text, entities)
            
            # Set result fields
            result.success = True
            result.text_content = text
            result.metadata = metadata
            result.entities = entities
            result.tags = tags
            result.confidence = 0.9 if text else 0.3
            
            return result
            
        except Exception as e:
            result.success = False
            result.error_message = f"Error processing Word document: {str(e)}"
            return result
    
    async def _extract_text_and_metadata(self, file_path: str) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text and metadata from a Word document.
        
        Args:
            file_path: Path to the Word document
            
        Returns:
            Tuple[str, Dict[str, Any]]: Extracted text and metadata
        """
        # Run in a thread to avoid blocking the event loop
        def extract():
            _, ext = os.path.splitext(file_path)
            
            # Initialize metadata
            metadata = {
                'file_type': ext
            }
            
            # Process based on file type
            if ext.lower() == '.docx':
                # Open the document
                doc = docx.Document(file_path)
                
                # Extract metadata
                core_properties = doc.core_properties
                if core_properties:
                    if core_properties.title:
                        metadata['title'] = core_properties.title
                    if core_properties.author:
                        metadata['author'] = core_properties.author
                    if core_properties.created:
                        metadata['created'] = core_properties.created.isoformat()
                    if core_properties.modified:
                        metadata['modified'] = core_properties.modified.isoformat()
                
                # Extract text
                paragraphs = []
                for para in doc.paragraphs:
                    if para.text:
                        paragraphs.append(para.text)
                
                # Add table content
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            if cell.text:
                                row_text.append(cell.text)
                        if row_text:
                            paragraphs.append(" | ".join(row_text))
                
                text = "\n".join(paragraphs)
                
                # Add document statistics
                metadata['paragraph_count'] = len(doc.paragraphs)
                metadata['table_count'] = len(doc.tables)
                
                return text, metadata
            else:
                # For .doc files, we can't extract content directly
                # Return empty text and basic metadata
                return "", metadata
        
        # Run in a thread pool
        return await asyncio.to_thread(extract)
    
    async def _extract_entities(self, text: str) -> Dict[str, Any]:
        """
        Extract entities from text.
        
        Args:
            text: Text to extract entities from
            
        Returns:
            Dict[str, Any]: Extracted entities
        """
        entities = {}
        
        # Extract invoice-related entities
        if any(keyword in text.lower() for keyword in ['invoice', 'faktura', 'rachunek']):
            # Extract invoice number
            invoice_number_match = re.search(r'(?:invoice|faktura|rachunek)[^\d]*(\d+[-/\w]*)', text.lower())
            if invoice_number_match:
                entities['invoice_number'] = invoice_number_match.group(1)
            
            # Extract amount
            amount_match = re.search(r'(?:amount|kwota|suma)[^\d]*(\d+[.,]\d+)', text.lower())
            if amount_match:
                entities['amount'] = amount_match.group(1)
            
            # Extract date
            date_match = re.search(r'(?:date|data)[^\d]*(\d{1,2}[-./]\d{1,2}[-./]\d{2,4})', text.lower())
            if date_match:
                entities['date'] = date_match.group(1)
        
        # Extract contract-related entities
        if any(keyword in text.lower() for keyword in ['contract', 'agreement', 'umowa']):
            # Extract contract number
            contract_number_match = re.search(r'(?:contract|agreement|umowa)[^\d]*(\d+[-/\w]*)', text.lower())
            if contract_number_match:
                entities['contract_number'] = contract_number_match.group(1)
            
            # Extract parties
            parties_match = re.search(r'(?:between|pomiędzy)\s+([^,]+),?\s+(?:and|i)\s+([^,\.]+)', text.lower())
            if parties_match:
                entities['party1'] = parties_match.group(1).strip()
                entities['party2'] = parties_match.group(2).strip()
        
        # Extract contact information
        email_matches = re.findall(r'[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[a-zA-Z]{2,}', text)
        if email_matches:
            entities['emails'] = email_matches
        
        phone_matches = re.findall(r'(?:\+\d{1,3}[-\s]?)?\(?\d{3}\)?[-\s]?\d{3}[-\s]?\d{3,4}', text)
        if phone_matches:
            entities['phones'] = phone_matches
        
        return entities
    
    async def _determine_tags(self, text: str, entities: Dict[str, Any]) -> List[str]:
        """
        Determine tags based on content.
        
        Args:
            text: Extracted text
            entities: Extracted entities
            
        Returns:
            List[str]: Tags
        """
        tags = []
        
        # Check for document type
        if any(keyword in text.lower() for keyword in ['invoice', 'faktura', 'rachunek']):
            tags.append('invoice')
        
        if any(keyword in text.lower() for keyword in ['contract', 'agreement', 'umowa']):
            tags.append('contract')
        
        if any(keyword in text.lower() for keyword in ['report', 'raport', 'sprawozdanie']):
            tags.append('report')
        
        if any(keyword in text.lower() for keyword in ['offer', 'proposal', 'oferta']):
            tags.append('offer')
        
        if any(keyword in text.lower() for keyword in ['letter', 'list']):
            tags.append('letter')
        
        # Check for HVAC-specific content
        if any(keyword in text.lower() for keyword in ['hvac', 'heating', 'ventilation', 'air conditioning', 'klimatyzacja', 'ogrzewanie', 'wentylacja']):
            tags.append('hvac')
        
        return tags